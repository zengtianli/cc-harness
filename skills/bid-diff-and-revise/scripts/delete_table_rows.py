#!/usr/bin/env python3
"""删除 docx 中指定表的指定行范围（track-changes 不支持表行删除时用）。

Usage:
  python3 delete_table_rows.py --docx 修改稿.docx --table-index 6 --rows 8:15 \
    --expected-first-col "序号,1,2,3,4,5,6,7" \
    --expected-residue "自然资源集约利用"
"""
import argparse
import sys
from pathlib import Path

from docx import Document


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--docx", required=True, help="目标 docx（原地修改）")
    ap.add_argument("--table-index", type=int, required=True, help="表格索引（0-based）")
    ap.add_argument("--rows", required=True, help="要删的行范围 FROM:TO（闭区间，0-based）")
    ap.add_argument("--expected-first-col", default="",
                    help="删除前保留行第一列期望值（逗号分隔），用于安全校验；留空则跳过校验")
    ap.add_argument("--expected-residue", default="",
                    help="要删的第一行第二列应含此关键字，确保删对了")
    ap.add_argument("--expected-last", default="",
                    help="删后末行第二列期望值，用于校验删对了")
    args = ap.parse_args()

    docx_path = Path(args.docx)
    doc = Document(str(docx_path))

    if args.table_index >= len(doc.tables):
        print(f"❌ table-index 超出范围（文档只有 {len(doc.tables)} 个表）", file=sys.stderr)
        sys.exit(1)

    tbl = doc.tables[args.table_index]
    from_idx, to_idx = map(int, args.rows.split(":"))

    print(f"删除前 T{args.table_index} 行数: {len(tbl.rows)}")
    print(f"计划删除: R{from_idx}-R{to_idx}（共 {to_idx-from_idx+1} 行）")

    # 安全校验 1：保留行的第一列
    if args.expected_first_col:
        expected = [x.strip() for x in args.expected_first_col.split(",")]
        actual = [tbl.rows[i].cells[0].text.strip() for i in range(min(from_idx, len(tbl.rows)))]
        if actual != expected:
            print(f"❌ 前 {from_idx} 行第一列与期望不一致：\n  期望 {expected}\n  实际 {actual}", file=sys.stderr)
            sys.exit(1)
        print("✓ 保留行结构校验通过")

    # 安全校验 2：要删的起始行应含关键字
    if args.expected_residue:
        if len(tbl.rows[from_idx].cells) >= 2:
            first_cell = tbl.rows[from_idx].cells[0].text
            second_cell = tbl.rows[from_idx].cells[1].text
            if args.expected_residue not in (first_cell + second_cell):
                # 检查其它被删行中是否至少有一行含关键字
                found = False
                for i in range(from_idx, min(to_idx + 1, len(tbl.rows))):
                    if args.expected_residue in " ".join(c.text for c in tbl.rows[i].cells):
                        found = True
                        break
                if not found:
                    print(f"❌ 被删行中未找到关键字「{args.expected_residue}」，可能删错行", file=sys.stderr)
                    sys.exit(1)
        print(f"✓ 被删行包含期望关键字「{args.expected_residue}」")

    # 倒序删除
    for i in range(to_idx, from_idx - 1, -1):
        if i < len(tbl.rows):
            row = tbl.rows[i]
            row._element.getparent().remove(row._element)

    doc.save(str(docx_path))

    # 验证
    doc2 = Document(str(docx_path))
    tbl2 = doc2.tables[args.table_index]
    print(f"删除后行数: {len(tbl2.rows)}")
    if args.expected_last and len(tbl2.rows) >= 1:
        last = tbl2.rows[-1]
        if len(last.cells) >= 2:
            actual_last = last.cells[1].text.strip()
            if args.expected_last and args.expected_last not in actual_last:
                print(f"⚠️ 末行第二列「{actual_last}」不含期望「{args.expected_last}」", file=sys.stderr)
            else:
                print(f"✓ 末行校验通过: {actual_last[:40]}")
    print(f"OK -> {docx_path}")


if __name__ == "__main__":
    main()
